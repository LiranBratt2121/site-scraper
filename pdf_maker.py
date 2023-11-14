from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Gamma Course', 0)


def add_page(Text: str, imgPath: str) -> None:  # Heading to be added
    # document.add_heading(Heading, level=1)
    document.add_paragraph(Text)
    if imgPath:
        document.add_picture(imgPath, width=Inches(1.25))


with open("files/yes", "r", encoding="utf-8") as file:
    w = file.read()


# document.add_picture('AGT.png', width=Inches(1.25))
add_page(w, "img/AGT.jpg")


document.add_page_break()

document.save('full.docx')
